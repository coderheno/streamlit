# app.py
"""
Streamlit app to display the full contents of a DOCX file (paragraphs + tables).
- Automatically uses /mnt/data/BCA301-4 Updated.docx if present.
- Also allows users to upload a .docx file.
- If python-docx is not installed, shows install instructions and allows paste/upload of plain text as fallback.

Run:
    pip install streamlit python-docx
    streamlit run app.py
"""

import streamlit as st
import re
from pathlib import Path
import io

st.set_page_config(page_title="BCA301-4 Course Plan Viewer", layout="wide")

st.title("BCA301-4 — Course Plan Viewer")
st.caption("Displays full contents of a DOCX (paragraphs & tables).")

# Try to import python-docx (Document). If not present, we keep going but show instructions.
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# -------------------------
# Utilities for DOCX parsing
# -------------------------
def read_docx_paragraphs(path_or_bytes) -> list:
    """
    Returns a list of paragraph texts in order from the docx.
    Accepts a Path-like or bytes/BytesIO.
    """
    doc = Document(path_or_bytes)
    paras = []
    for p in doc.paragraphs:
        paras.append(p.text)
    return paras

def read_docx_tables(path_or_bytes) -> list:
    """
    Returns a list of tables, where each table is a list of rows,
    and each row is a list of cell text.
    """
    doc = Document(path_or_bytes)
    tables = []
    for table in doc.tables:
        rows = []
        for r in table.rows:
            cells = [c.text for c in r.cells]
            rows.append(cells)
        tables.append(rows)
    return tables

def extract_headings_and_sections(paragraphs: list) -> dict:
    """
    Heuristic: split text into sections based on lines that start with 'SECTION' (common in your syllabus).
    Returns dict: {section_heading: full_text_in_section}
    If no SECTION headings found, return {"Full Document": full_text}
    """
    text = "\n\n".join(paragraphs)
    SECTION_PATTERN = re.compile(r"(SECTION\s+[IVXLC]+)", flags=re.IGNORECASE)
    parts = SECTION_PATTERN.split(text)
    if len(parts) == 1:
        return {"Full Document": text}
    sections = {}
    i = 1
    while i < len(parts):
        header = parts[i].strip()
        content = parts[i + 1].strip() if (i + 1) < len(parts) else ""
        sections[header] = content
        i += 2
    return sections

# -------------------------
# Load docx: automatic file + upload
# -------------------------
DEFAULT_DOCX_PATH = Path("/mnt/data/BCA301-4 Updated.docx")

uploaded_docx = st.sidebar.file_uploader("Upload a .docx file (optional)", type=["docx"])

use_default = False
if DEFAULT_DOCX_PATH.exists() and uploaded_docx is None:
    use_default = st.sidebar.checkbox(f"Use existing file: {DEFAULT_DOCX_PATH.name}", value=True)
else:
    # If default not present, don't show check; user can upload
    use_default = False

docx_bytes = None
docx_source_name = None

if use_default and DEFAULT_DOCX_PATH.exists():
    docx_bytes = str(DEFAULT_DOCX_PATH)  # path; our read functions accept path string
    docx_source_name = DEFAULT_DOCX_PATH.name
elif uploaded_docx is not None:
    # streamlit gives a UploadedFile which behaves like BytesIO
    docx_bytes = io.BytesIO(uploaded_docx.read())
    docx_source_name = uploaded_docx.name

# -------------------------
# Fallback text input (when python-docx missing or user wants to paste)
# -------------------------
st.sidebar.markdown("---")
st.sidebar.write("Alternate input (paste text or upload .txt/.md)")

paste_text = st.sidebar.text_area("Paste course content (plain text)", height=120)
txt_upload = st.sidebar.file_uploader("Or upload .txt/.md", type=["txt", "md"])

txt_uploaded_content = None
if txt_upload is not None:
    txt_uploaded_content = txt_upload.read().decode("utf-8")

# -------------------------
# Main display
# -------------------------
col_main, col_right = st.columns([3, 1])

with col_main:
    st.header("Document Contents")

    if docx_bytes and HAS_DOCX:
        try:
            # Read paragraphs and tables
            paras = read_docx_paragraphs(docx_bytes)
            tables = read_docx_tables(docx_bytes)

            st.success(f"Loaded DOCX: {docx_source_name}")

            # Extract and show sections (based on 'SECTION' headings)
            sections = extract_headings_and_sections(paras)
            # Sidebar navigation for sections
            st.sidebar.header("Sections")
            section_keys = list(sections.keys())
            selected_section = st.sidebar.selectbox("Select section to view", section_keys, index=0)

            # Search inside document (global)
            search_query = st.sidebar.text_input("Search whole document")
            if search_query:
                # naive search through paragraphs
                matches = []
                for i, p in enumerate(paras):
                    if search_query.lower() in p.lower():
                        matches.append((i + 1, p.strip()))
                st.info(f"Found {len(matches)} paragraph matches for '{search_query}'.")
                for idx, text in matches[:200]:
                    st.write(f"Paragraph {idx}: {text}")

            # Display the selected section's text (split into blocks)
            st.subheader(selected_section)
            section_text = sections.get(selected_section, "")
            if section_text.strip() == "":
                st.write("_(no textual content detected in this section)_")
            else:
                blocks = [b for b in re.split(r"\n{2,}", section_text) if b.strip()]
                for i, b in enumerate(blocks):
                    if len(b) > 400:
                        with st.expander(f"Part {i+1}"):
                            st.write(b)
                    else:
                        st.write(b)

            # Show tables (if any) after the textual contents
            if tables:
                st.markdown("---")
                st.subheader("Tables found in document")
                for t_idx, table in enumerate(tables, start=1):
                    st.markdown(f"**Table {t_idx}**")
                    # Render as st.table using list of dicts if header-like row detected
                    # Heuristic: if first row looks like headers (no empty cells), use it
                    rows = table
                    if not rows:
                        continue
                    # Use the first row as headers if all unique
                    headers = rows[0]
                    # build list of dicts for dataframe
                    data = []
                    for r in rows[1:]:
                        # pad shorter rows
                        row_cells = r + [""] * (len(headers) - len(r))
                        data.append({headers[j] if headers[j] else f"Col {j+1}": row_cells[j] for j in range(len(headers))})
                    if data:
                        st.table(data)
                    else:
                        # only single-row table, show as single-row table
                        st.write(headers)

            # Provide download of full extracted text
            full_text = "\n\n".join(paras)
            st.download_button("Download extracted text (.txt)", full_text, file_name="BCA301-4_extracted.txt", mime="text/plain")

        except Exception as e:
            st.error(f"Error parsing DOCX: {e}")
            st.write("You can paste the document text in the sidebar or upload a .txt/.md file as fallback.")
    else:
        # Either no docx uploaded/selected, or python-docx missing
        if docx_bytes and not HAS_DOCX:
            st.warning("python-docx is not installed in the environment. Install it to enable .docx parsing.")
            st.info("Install by running: pip install python-docx")
            st.write("After installing, re-run the Streamlit app.")
            st.markdown("---")

        # Prefer pasted text, then txt upload, else show message
        input_text = None
        if paste_text and paste_text.strip():
            input_text = paste_text
            st.success("Showing pasted text")
        elif txt_uploaded_content:
            input_text = txt_uploaded_content
            st.success(f"Showing uploaded text: {txt_upload.name}")
        else:
            st.info("No DOCX parsed. Either upload a .docx (and have python-docx installed) or paste/upload plain text in the sidebar.")
            input_text = ""

        if input_text:
            # Try to split into SECTIONs if present
            SECTION_PATTERN = re.compile(r"(SECTION\s+[IVXLC]+)", flags=re.IGNORECASE)
            parts = SECTION_PATTERN.split(input_text)
            if len(parts) == 1:
                st.write(input_text)
            else:
                sections = {}
                i = 1
                while i < len(parts):
                    header = parts[i].strip()
                    content = parts[i+1].strip() if (i+1) < len(parts) else ""
                    sections[header] = content
                    i += 2
                st.sidebar.header("Sections")
                selected_section = st.sidebar.selectbox("Select section to view", list(sections.keys()), index=0)
                st.subheader(selected_section)
                st.write(sections[selected_section])
            # download button for pasted/ uploaded text
            if input_text.strip():
                st.download_button("Download shown text (.txt)", input_text, file_name="BCA301-4_shown.txt", mime="text/plain")

with col_right:
    st.markdown("### Quick actions")
    if HAS_DOCX:
        st.write("- `python-docx` is available.")
    else:
        st.warning("- `python-docx` not installed.")
        st.write("Install with:")
        st.code("pip install python-docx")
    st.markdown("---")
    st.write("Tips:")
    st.write("• If you have `/mnt/data/BCA301-4 Updated.docx` on the server, enable the checkbox to load it.")
    st.write("• You can upload a .docx file or paste the content as fallback.")
    st.write("• Tables are displayed as Streamlit tables when possible.")

st.markdown("---")
st.caption("This app attempts to show the full contents of your DOCX (paragraphs & tables).")
